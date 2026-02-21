"""
routes/export.py — .docx export from HTML content
Requires: python-docx, beautifulsoup4
"""

import re
import uuid
import tempfile
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from middleware.auth import get_current_user

router = APIRouter(prefix="/notebooks", tags=["export"])


class DocxExportRequest(BaseModel):
    title: str
    html: str


@router.post("/{nb_id}/export/docx")
async def export_docx(
    nb_id: str,
    body: DocxExportRequest,
    username: str = Depends(get_current_user),
):
    """Convert notebook HTML to a .docx file and return a download URL."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from bs4 import BeautifulSoup, NavigableString, Tag
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx and beautifulsoup4 are required. Run: pip install python-docx beautifulsoup4"
        )

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title_para = doc.add_heading(body.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    soup = BeautifulSoup(body.html, 'html.parser')

    def add_run_with_style(para, text, el):
        run = para.add_run(text)
        # Bold
        if el.find_parent(['b', 'strong']) or el.name in ['b', 'strong']:
            run.bold = True
        # Italic
        if el.find_parent(['i', 'em']) or el.name in ['i', 'em']:
            run.italic = True
        # Underline
        if el.find_parent('u') or el.name == 'u':
            run.underline = True
        # Strikethrough
        if el.find_parent('s') or el.name == 's':
            run.font.strike = True
        # Font size from style attribute
        style_attr = el.get('style', '') if hasattr(el, 'get') else ''
        size_match = re.search(r'font-size:\s*([\d.]+)px', style_attr)
        if size_match:
            run.font.size = Pt(float(size_match.group(1)) * 0.75)
        # Font color
        color_match = re.search(r'color:\s*(#[0-9a-fA-F]{6})', style_attr)
        if color_match:
            hex_c = color_match.group(1).lstrip('#')
            run.font.color.rgb = RGBColor(int(hex_c[0:2],16), int(hex_c[2:4],16), int(hex_c[4:6],16))
        return run

    def process_element(el, parent_para=None):
        if isinstance(el, NavigableString):
            text = str(el)
            if text.strip() and parent_para:
                add_run_with_style(parent_para, text, el)
            return

        tag = el.name if hasattr(el, 'name') else None
        if tag is None:
            return

        # Block elements
        if tag in ['h1', 'h2', 'h3']:
            level = int(tag[1])
            p = doc.add_heading('', level=level)
            for child in el.children:
                process_element(child, p)
        elif tag == 'p':
            p = doc.add_paragraph()
            for child in el.children:
                process_element(child, p)
        elif tag == 'blockquote':
            p = doc.add_paragraph(style='Quote')
            for child in el.children:
                process_element(child, p)
        elif tag == 'pre':
            code_text = el.get_text()
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif tag == 'hr':
            doc.add_paragraph('─' * 50)
        elif tag in ['ul', 'ol']:
            style_name = 'List Bullet' if tag == 'ul' else 'List Number'
            for li in el.find_all('li', recursive=False):
                p = doc.add_paragraph(style=style_name)
                for child in li.children:
                    process_element(child, p)
        elif tag == 'table':
            rows = el.find_all('tr')
            if not rows:
                return
            cols = max(len(r.find_all(['td','th'])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            for ri, row in enumerate(rows):
                cells = row.find_all(['td','th'])
                for ci, cell in enumerate(cells):
                    if ci < cols:
                        tc = table.cell(ri, ci)
                        tc.text = cell.get_text()
                        if cell.name == 'th':
                            for run in tc.paragraphs[0].runs:
                                run.bold = True
        elif tag in ['br']:
            if parent_para:
                parent_para.add_run('\n')
        elif tag in ['div', 'section', 'article']:
            # Treat as block
            has_block = any(hasattr(c,'name') and c.name in ['p','h1','h2','h3','ul','ol','table','pre','blockquote'] for c in el.children)
            if has_block or parent_para is None:
                p = doc.add_paragraph()
                for child in el.children:
                    process_element(child, p)
            else:
                for child in el.children:
                    process_element(child, parent_para)
        else:
            # Inline element — recurse into parent para
            for child in el.children:
                process_element(child, parent_para)

    # Process all top-level children
    for el in soup.children:
        process_element(el)

    # Save to temp file
    tmp_dir  = Path(tempfile.gettempdir()) / 'workspace_exports'
    tmp_dir.mkdir(exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    filepath = tmp_dir / filename
    doc.save(str(filepath))

    return {"download_url": f"/notebooks/{nb_id}/export/download/{filename}"}


@router.get("/{nb_id}/export/download/{filename}")
async def download_export(
    nb_id: str,
    filename: str,
    username: str = Depends(get_current_user),
):
    """Serve a previously generated export file."""
    # Sanitize filename — only allow hex + .docx
    if not re.match(r'^[a-f0-9]{32}\.docx$', filename):
        raise HTTPException(status_code=400, detail="Invalid filename.")
    filepath = Path(tempfile.gettempdir()) / 'workspace_exports' / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Export file not found or expired.")
    return FileResponse(
        path=str(filepath),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename,
    )
